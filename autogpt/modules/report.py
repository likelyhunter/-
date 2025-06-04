from docx import Document

def generate_report(target, recon, scan, exploit, persistence):
    doc = Document()
    doc.add_heading('渗透测试报告', 0)
    doc.add_paragraph(f"目标系统: {target}")

    doc.add_heading('信息收集', level=1)
    doc.add_paragraph(recon)

    doc.add_heading('漏洞扫描', level=1)
    doc.add_paragraph(scan)

    doc.add_heading('漏洞利用', level=1)
    doc.add_paragraph(exploit)

    doc.add_heading('权限维持', level=1)
    doc.add_paragraph(persistence)

    doc.add_heading('总结与建议', level=1)
    doc.add_paragraph("报告由 GPT 自动生成，请进一步审核。")

    doc.save(f"Pentest_Report_{target}.docx")

#功能：自动生成 Word 报告，结构清晰、模块化
#可与 GPT 结合，让 GPT 生成“渗透路径总结”、“安全建议”等段落

